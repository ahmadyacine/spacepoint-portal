import os
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import docx
import platform
import subprocess
from docx2pdf import convert as docx2pdf_convert
from datetime import datetime
from app.core.config import settings

def convert_to_pdf(docx_path: str, pdf_path: str):
    """
    Converts DOCX to PDF. Works on Windows (using docx2pdf) and Linux (using libreoffice).
    """
    if platform.system() == "Windows":
        try:
            import pythoncom
            pythoncom.CoInitialize()
        except ImportError:
            print("[email_service] pythoncom/pywin32 not installed, skipping CoInitialize.")
        try:
            docx2pdf_convert(docx_path, pdf_path)
        except Exception as e:
            print(f"[email_service] docx2pdf failed: {e}. Trying libreoffice fallback...")
            try:
                # Try libreoffice on windows if docx2pdf fails
                subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path], check=True)
            except:
                print(f"[email_service] All PDF conversion methods failed on Windows.")
                raise e
    else:
        # Linux / Other - use libreoffice
        try:
            # Note: outdir must be absolute
            outdir = os.path.abspath(os.path.dirname(pdf_path))
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', outdir, docx_path], check=True)
            # Libreoffice saves it as original_name.pdf in the outdir.
            # If the filenames differ, rename it.
            expected_pdf = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if expected_pdf != pdf_path and os.path.exists(expected_pdf):
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                os.rename(expected_pdf, pdf_path)
        except Exception as e:
            print(f"[email_service] Linux PDF conversion failed (is libreoffice installed?): {e}")
            raise e

def generate_contract(instructor_name: str, living_area: str) -> tuple[str, str]:
    """
    Generates a personalized contract from the template.
    Returns (docx_path, pdf_path).
    """
    # Calculate relative paths to project root
    # Current file: .../backend/app/services/email_service.py
    # Root is up 3 levels
    current_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.abspath(os.path.join(current_dir, "..", "..", ".."))
    
    template_path = os.path.join(project_root, "SPACE.-FC-AGREEMENTLETTER-EN - November.docx")
    output_dir = os.path.join(project_root, "backend", "app", "uploads", "contracts")
    
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    now = datetime.now()
    month_name = now.strftime("%B")
    current_date_str = now.strftime("%d %B %Y")
    
    # Sanitize instructor name for filename
    safe_name = "".join(x for x in instructor_name if x.isalnum() or x in " -_").strip().replace(" ", "")
    base_filename = f"{safe_name}_SPACE.-FC-AGREEMENTLETTER-EN - {month_name}"
    
    docx_output = os.path.join(output_dir, f"{base_filename}.docx")
    pdf_output = os.path.join(output_dir, f"{base_filename}.pdf")
    
    doc = docx.Document(template_path)
    
    replacements = {
        "[Name of Instructor]": instructor_name,
        "Al Ain, Abu Dhabi": living_area,
        "1 November 2025": current_date_str
    }
    
    def apply_bold_replacements(paragraphs):
        for p in paragraphs:
            # We process each placeholder independently
            for placeholder, replacement in replacements.items():
                if placeholder in p.text:
                    # Check if it's entirely in one run
                    found = False
                    for run in p.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)
                            run.bold = True
                            found = True
                    
                    # If split across runs, we have to rebuild the paragraph text
                    # to ensure the replacement happens.
                    if not found:
                        current_text = p.text
                        new_text = current_text.replace(placeholder, replacement)
                        
                        # Find indices for bolding
                        start_idx = current_text.find(placeholder)
                        before = current_text[:start_idx]
                        after = current_text[start_idx + len(placeholder):]
                        
                        # Clear old runs and add new ones
                        # p.text = "" clears all runs in python-docx
                        p.text = "" 
                        p.add_run(before)
                        r_mid = p.add_run(replacement)
                        r_mid.bold = True
                        p.add_run(after)

    apply_bold_replacements(doc.paragraphs)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                apply_bold_replacements(cell.paragraphs)
                            
    doc.save(docx_output)
    
    # Convert to PDF
    try:
        convert_to_pdf(docx_output, pdf_output)
    except Exception as e:
        print(f"[email_service] Error converting to PDF: {str(e)}")
        # If conversion fails, we return None for PDF path
        pdf_output = None
        
    return docx_output, pdf_output

def send_approval_credentials_email(to_email: str, name: str, temp_password: str, living_area: str = "Unknown") -> tuple[bool, str]:
    """
    Sends an HTML email with temporary login credentials and attached PDF contract.
    Returns (True/False, pdf_contract_path).
    """
    smtp_host = settings.SMTP_HOST
    smtp_port = settings.SMTP_PORT
    smtp_user = settings.SMTP_USER
    smtp_password = settings.SMTP_PASSWORD

    print(f"[email_service] Sending to: {to_email}")
    print(f"[email_service] SMTP_USER: {smtp_user}")
    print(f"[email_service] SMTP_PASSWORD (first 4): {smtp_password[:4] if smtp_password else 'NOT SET'}...")

    msg = MIMEMultipart("alternative")
    msg["Subject"] = "SpacePoint Instructor Application Approved"
    msg["From"] = f"SpacePoint <{smtp_user}>"
    msg["To"] = to_email

    LOGO_URL = "https://spacepoint.ae/wp-content/uploads/2023/12/SpacePoint-Purple-Logo-e1766489433825-1024x261.png"

    html_body = f"""
<!DOCTYPE html>
<html>
<body style="margin:0;padding:0;background-color:#f3f4f6;font-family:system-ui,Arial,sans-serif;">
  <table width="100%" cellpadding="0" cellspacing="0" style="padding:32px 0;background:#f3f4f6;">
    <tr><td align="center">
      <table width="600" cellpadding="0" cellspacing="0"
             style="background:#ffffff;border-radius:16px;overflow:hidden;
                    box-shadow:0 4px 24px rgba(36,17,52,0.12);">

        <!-- Header with SpacePoint logo -->
        <tr>
          <td style="background:linear-gradient(135deg,#241134,#653f84);padding:28px 32px;">
            <img src="{LOGO_URL}" height="44" alt="SpacePoint" style="display:block;">
          </td>
        </tr>

        <!-- Body -->
        <tr>
          <td style="padding:36px 32px;">

            <h2 style="margin:0 0 6px 0;font-size:20px;font-weight:700;color:#1a1135;">
              Application Approved &#10003;
            </h2>
            <p style="margin:0 0 24px 0;font-size:12px;font-weight:700;color:#653f84;
                       text-transform:uppercase;letter-spacing:0.1em;">
              Instructor Scholarship Programme
            </p>

            <p style="margin:0 0 12px 0;font-size:15px;color:#374151;line-height:1.7;">
              Hello <strong>{name}</strong>,
            </p>
            <p style="margin:0 0 24px 0;font-size:15px;color:#374151;line-height:1.7;">
              Congratulations! Your
              <strong>SpacePoint Instructor Scholarship Application</strong>
              has been reviewed and
              <strong style="color:#653f84;">approved</strong>.
              You now have access to the SpacePoint Instructor Portal.
            </p>

            <!-- Credentials box -->
            <div style="background:#faf8ff;border:1px solid #ddd6fe;border-radius:12px;
                         padding:22px 26px;margin:0 0 24px 0;">
              <p style="margin:0 0 16px 0;font-size:11px;font-weight:700;
                         letter-spacing:0.18em;color:#653f84;text-transform:uppercase;">
                Your Login Credentials
              </p>
              <table width="100%" cellpadding="0" cellspacing="0">
                <tr>
                  <td style="padding:8px 0;border-bottom:1px solid #ede9f7;">
                    <span style="font-size:11px;color:#9ca3af;text-transform:uppercase;
                                 letter-spacing:0.1em;">Email</span><br>
                    <span style="font-family:'Courier New',monospace;font-size:15px;
                                 color:#1a1135;font-weight:600;">{to_email}</span>
                  </td>
                </tr>
                <tr>
                  <td style="padding:14px 0 0 0;">
                    <span style="font-size:11px;color:#9ca3af;text-transform:uppercase;
                                 letter-spacing:0.1em;">Temporary Password</span><br>
                    <span style="font-family:'Courier New',monospace;font-size:20px;
                                 color:#653f84;font-weight:700;letter-spacing:0.05em;">{temp_password}</span>
                  </td>
                </tr>
              </table>
            </div>

            <!-- Warning note -->
            <div style="background:#fffbeb;border-left:4px solid #f59e0b;
                         border-radius:0 8px 8px 0;padding:14px 18px;margin:0 0 28px 0;">
              <p style="margin:0;font-size:13px;color:#92400e;line-height:1.6;">
                <strong>&#9888; Action Required:</strong>
                This is a temporary password. Please log in and set a personal
                password before accessing your instructor resources.
              </p>
            </div>

            <!-- CTA button -->
            <table cellpadding="0" cellspacing="0">
              <tr>
                <td style="background:linear-gradient(135deg,#241134,#653f84);border-radius:10px;">
                  <a href="http://localhost:8000"
                     style="display:inline-block;padding:13px 32px;font-size:14px;
                            font-weight:700;color:#ffffff;text-decoration:none;letter-spacing:0.04em;">
                    Log In to Instructor Portal &rarr;
                  </a>
                </td>
              </tr>
            </table>

            <p style="margin:24px 0 0 0;font-size:15px;color:#374151;line-height:1.7;">
              <strong>Note:</strong> We have attached your contract to this email. Please review it, sign it, and upload it back into the portal.
            </p>

            <p style="margin:28px 0 0 0;font-size:14px;color:#6b7280;line-height:1.7;">
              Best regards,<br>
              <strong style="color:#241134;">The SpacePoint Team</strong>
            </p>

          </td>
        </tr>

        <!-- Footer -->
        <tr>
          <td style="background:#f9fafb;padding:18px 32px;border-top:1px solid #ede9f7;">
            <p style="margin:0;font-size:12px;color:#9ca3af;text-align:center;line-height:1.7;">
              &copy; 2026 SpacePoint &nbsp;&middot;&nbsp; www.spacepoint.ae<br>
              <em>Do not reply to this email. This mailbox is not monitored.</em>
            </p>
          </td>
        </tr>

      </table>
    </td></tr>
  </table>
</body>
</html>
"""

    msg.attach(MIMEText(html_body, "html"))

    # Generate and attach contract
    docx_path, pdf_path = generate_contract(name, living_area)
    
    # We prefer the PDF version, but fallback to DOCX if PDF generation failed
    if pdf_path and os.path.exists(pdf_path):
        file_path = pdf_path
    elif docx_path and os.path.exists(docx_path):
        file_path = docx_path
        print(f"[email_service] Warning: PDF conversion failed, attaching DOCX instead: {docx_path}")
    else:
        file_path = None
    
    if file_path:
        try:
            with open(file_path, "rb") as attachment:
                part = MIMEBase("application", "octet-stream")
                part.set_payload(attachment.read())
                encoders.encode_base64(part)
                filename = os.path.basename(file_path)
                part.add_header(
                    "Content-Disposition",
                    f"attachment; filename= {filename}",
                )
                msg.attach(part)
        except Exception as e:
            print(f"[email_service] Failed to attach {file_path}: {str(e)}")
    else:
        print(f"[email_service] PDF contract not available for attachment. (docx exists at {docx_path})")

    try:
        server = smtplib.SMTP(smtp_host, smtp_port)
        server.starttls()
        server.login(smtp_user, smtp_password)
        server.sendmail(smtp_user, to_email, msg.as_string())
        server.quit()
        print(f"[email_service] Email sent successfully to {to_email}")
        return True, pdf_path
    except Exception as e:
        import traceback
        print(f"[email_service] Failed to send email to {to_email}: {str(e)}")
        traceback.print_exc()
        return False, pdf_path
