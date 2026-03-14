import os
import docx
from datetime import datetime

def test_replacement():
    template_path = r'c:\Users\ahmad yacine\Desktop\PortalV2\SPACE.-FC-AGREEMENTLETTER-EN - November.docx'
    output_docx = r'c:\Users\ahmad yacine\Desktop\PortalV2\backend\app\uploads\test_contract.docx'
    
    if not os.path.exists(r'c:\Users\ahmad yacine\Desktop\PortalV2\backend\app\uploads'):
        os.makedirs(r'c:\Users\ahmad yacine\Desktop\PortalV2\backend\app\uploads')

    doc = docx.Document(template_path)
    
    instructor_name = "John Doe"
    living_area = "Dubai, UAE"
    current_date = datetime.now().strftime("%d %B %Y")
    
    # Placeholders to replace
    replacements = {
        "[Name of Instructor]": instructor_name,
        "Al Ain, Abu Dhabi": living_area,
        "1 November 2025": current_date
    }
    
    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    # Also check tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in replacements.items():
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, replacement)
                            
    doc.save(output_docx)
    print(f"Saved to {output_docx}")

if __name__ == "__main__":
    test_replacement()
